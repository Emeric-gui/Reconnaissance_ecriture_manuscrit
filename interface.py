# on import les modules pour créer l'interface
import tkinter as tk
from tkinter import *
from tkinter import filedialog, ttk

import cv2
from PIL import Image, ImageTk
# on importe la librairie pour enregistrer dans un word
import docx

import os
from traitement import Traitement
from reseau import Classifier


class Fenetre:
    def __init__(self):
        self.__filename = None
        self.__window = tk.Tk()
        self.__window.geometry("500x500")  # Size of the window
        self.__window.title('Scribo')
        my_font1 = ('times', 14, 'bold')
        l1 = tk.Label(self.__window, text='Selectioner votre image', width=30, font=my_font1)
        l1.grid(row=1, column=1)  # on met notre texte en ligne 1 colonne 1
        b1 = tk.Button(self.__window, text='Upload Files', width=20, command=lambda: self.__upload_file())
        b1.grid(row=3, column=1)  # on met le bouton en dessous du texte soit en ligne 2 colonne 2

        self.__reseau = Classifier()

    def __upload_file(self):
        file_types = [('Images', '*.jpg;*.png;*.jpeg'),
                      ('All', '*.*')]  # type of files to select
        self.__filename = tk.filedialog.askopenfilename(title="Choisir une image", filetypes=file_types)
        img = None
        # on va ensuite afficher l'images choisie
        try:
            img = Image.open(self.__filename)
            img = img.resize((200, 200))  # new width & height
            img = ImageTk.PhotoImage(img)
            e1 = Label(image=img)
            e1.image = img
            e1.grid(row=4, column=1)
            b2 = tk.Button(self.__window, text='Transformer en Word', width=20, command=lambda: self.__imgToTxt())
            b2.grid(row=5, column=1)
        except AttributeError:
            print("Pas de fichier envoyé")



    def __imgToTxt(self):
        # on apply ici la fonction de traitement de l'image
        self.__applyTraitement(self.__filename)
        mydoc = docx.Document()  # on ouvre le word / le crée s'il existe pas
        mydoc = self.__feed_word(mydoc)
        champs_chemin = self.__filename.split("/")
        finChemin = champs_chemin[len(champs_chemin) - 1].split(".")[0]
        wordname = finChemin+".docx"
        path = "/"
        mydoc.save(wordname)
        letters = os.listdir("letters")
        for image in letters:
            os.remove("letters/" + image)
        e2 = Label(self.__window, text=wordname)
        e2.grid(row=6, column=1)
        print("Document fini")

    def activateWindow(self):
        self.__window.mainloop()  # Keep the window open

    def __applyTraitement(self, image):
        print("Image : {0}".format(image))
        cheminImage = image
        traitement = Traitement(cheminImage)
        traitement.pretraitement()

    def __feed_word(self, mydoc):
        # self.__reseau.test_predict()
        repertoire = "letters"
        if os.path.exists(repertoire):
            liste_images = os.listdir(repertoire)
            i = 1

            len_repertoire = len(liste_images)
            num_image = 0

            texte_ligne = ""
            ligne = 1
            mot = 1
            lettre = 1

            debut_ligne = True
            # in_word = False

            # Parcours de toutes les images et traitement sur chacune d'elles
            for str_image in liste_images:
                num_image += 1
                print("Image : {0}".format(str_image))

                # On récupère l'identifiant de l'image, permet de savoir sur quelle ligne, mot et lettre nous sommes
                temp_str_image = str_image.split(".")[0]
                str_image_split = temp_str_image.split("_")
                num_ligne = int(str_image_split[1])
                num_mot = int(str_image_split[2])

                if num_ligne != ligne:
                    texte_ligne += "."
                    mydoc.add_paragraph(texte_ligne)
                    texte_ligne = ""
                    ligne += 1
                    debut_ligne = True
                elif num_mot != mot:
                    texte_ligne += " "
                    mot += 1
                    debut_ligne = False
                    # in_word = False
                else:
                    debut_ligne = False
                    # in_word = True

                if num_image == 1:
                    debut_ligne = True

                image = cv2.imread(repertoire+"/"+str_image)
                image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                texte_ligne += self.__reseau.reality(image, debut_ligne)

                # Check si on est a la derniere image du dossier lettres, si ou i
                if num_image == len_repertoire:
                    mydoc.add_paragraph(texte_ligne)
        return mydoc
